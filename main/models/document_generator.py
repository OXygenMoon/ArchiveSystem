# models/document_generator.py

import os
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from config import ARCHIVE_DIR

def set_font(run, font_name='宋体', size_pt=12):
    """
    设置文档中 run 对象的字体和字号。
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)


def get_today_date_str():
    """
    获取当前日期并返回 "YYYY.MM.DD" 格式的字符串。
    """
    today = date.today()
    return today.strftime("%Y.%m.%d")

def create_document(student_name, sex, student_class, reason, level, department):
    """
    生成处分文档（.docx）并保存到 ARCHIVE_DIR 文件夹下。
    返回生成的文件名。
    """
    doc = Document()

    # 标题部分
    title = doc.add_heading('学 生 处 分\n')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_font(run, size_pt=22)  # 可以把标题字号加大

    content_template = (
        f"{student_name}(性别 {sex} )现就读我校{student_class}班。"
        f"该同学由于{reason}，造成不良影响。该同学的行为违反了《中学生日常行为规范》"
        f"及嘉善技师学院（筹）·嘉善县中等专业学校的相关校规，经{department}和班主任研究，"
        f"报学校政教处审批，决定给予{{level}}处分。{{tail}}\n"
    )

    # 根据处分级别动态拼接尾部内容
    if level != '留校察看':
        content_text = content_template.format(
            level=level,
            tail="若该生仍不听从老师教诲，则由家长或监护人领回家继续教育，直到接受老师教育改正为止。"
        )
    else:
        content_text = content_template.format(
            level=level,
            tail="若该生仍不听从老师教诲，则自愿退学。"
        )

    paragraph = doc.add_paragraph(content_text)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    for run in paragraph.runs:
        set_font(run)

    today = date.today()
    info_text = (
        f"嘉善技师学院（筹）·嘉善县中等专业学校{department}\n"
        f"{today.year}年{today.month}月{today.day}日\n"
    )
    info = doc.add_paragraph(info_text)
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in info.runs:
        set_font(run)

    signed_text = "学生签名：\n\n家长签名：\n\n班主任签名："
    signed = doc.add_paragraph(signed_text)
    for run in signed.runs:
        set_font(run)

    # 文件名
    filename = f"{get_today_date_str()} {student_name} {student_class} {level}.docx"
    # 确保存档目录存在
    os.makedirs(ARCHIVE_DIR, exist_ok=True)
    doc_path = os.path.join(ARCHIVE_DIR, filename)
    doc.save(doc_path)

    return filename

