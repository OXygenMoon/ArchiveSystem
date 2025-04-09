import os
import json
import redis
import math
from flask import Flask, render_template, request, jsonify, redirect, url_for, session, send_file
# from flask_login import LoginManager, UserMixin, login_user, login_required, current_user, logout_user
# from flask_sqlalchemy import SQLAlchemy
from flask_session import Session
import pandas as pd
from sqlalchemy.pool import  reset_none
import random
import datetime
from io import BytesIO
import markdown

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING

# 数据绘图
from pyecharts import options as opts
from pyecharts.charts import Bar

from config import (
    SECRET_KEY, USER_ROLES, DEPARTMENTS,
    BASE_DIR, ARCHIVE_CSV, UPLOAD_FOLDER, ALLOWED_UPLOAD_EXTENSIONS, WEIJI_TYPES,
    WEIJI_LEVELS,  # 违纪等级列表
    DATA_FILE_JINGYI, DATA_FILE_ZHIZAO  # 经艺系 / 制造系 违纪表格路径
)

from utils.common import (
    load_data, load_session
)

from utils.decorators import required_role, check_redis_session

app = Flask(__name__)
app.secret_key = SECRET_KEY  # 使用 config.py 中的 SECRET_KEY

# 配置 Redis
app.config['SESSION_TYPE'] = 'redis'
app.config['SESSION_REDIS'] = redis.Redis(
    host='localhost',  # Redis 服务器地址
    port=6379,
    db=0,
    # password='your_redis_password'  # 如果设置了密码
)
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_USE_SIGNER'] = True  # 对 Session ID 签名

# 初始化 Session
Session(app)

# 日志
logger = app.logger


# 上传目录
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)


@app.after_request
def add_header(response):
    """
    在每个响应中添加 Cache-Control 头，防止缓存。
    """
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response

@app.route('/')
def index():
    """
    主页路由，若已登录则跳转到 index，否则渲染 home.html。
    """
    if session.get('logged_in'):
        return redirect(url_for('home'))
    welcome = '欢迎来到德育管理平台'
    return render_template('index.html', msg=welcome)


@app.route('/logout')
def logout():
    """
    退出操作: 清除session, 新建session, 重定向回登录页面
    """
    session.clear()  # 1. 清除会话数据
    session.regenerate()  # 2. (推荐) 刷新会话 ID
    return redirect(url_for('index'))


@app.route('/login', methods=['POST', 'GET'])
def login():
    # POST
    if request.method == 'POST':
        username = request.form['name']
        password = request.form['password']

        with open('data/user.json', 'r', encoding='utf-8') as f:
            user = json.loads(f.read())

        if username in user:
            if password == user[username]['password']:  # 使用加密验证

                # 加载登录用户信息至 session
                # 用户名 / 角色 / 系部 / 班级 / 姓名

                session['logged_in'] = True
                session['username'] = user[username]['username']
                session['role'] = user[username]['role']
                session['department'] = user[username]['department']
                session['class'] = user[username]['class']
                session['truename'] = user[username]['truename']
                session.permanent = True
                return redirect(url_for('home'))
            else:
                logger.error(f"用户名密码错误: {username}")
                return render_template('login.html', prompt='密码错误')
        else:
            logger.error(f"用户名不存在: {username}")
            return render_template('login.html', prompt='账号不存在')

    # GET
    return render_template('login.html', msg='登录')


@app.route('/register', methods=['POST', 'GET'])
def register():
    pass


import random
@app.route('/home')
@check_redis_session  # 重启后重定向login
def home():
    # 1. 加载 session
    response_data = load_session()  # 从 session 加载数据
    quotes = quotes = [
    {"quote_text": "德育是教育之首，是培养全面发展的人才的根本保证。", "author": ""},
    {"quote_text": "教育的根本任务在于立德树人。", "author": "习近平"},
    {"quote_text": "教书育人，育人为本；德智体美，德育为先。", "author": ""},
    {"quote_text": "师者，所以传道受业解惑也。", "author": "韩愈"},
    {"quote_text": "爱是教育的灵魂，没有爱就没有教育。", "author": "苏霍姆林斯基"},
    {"quote_text": "教育植根于爱。", "author": "福禄贝尔"},
    {"quote_text": "没有德育，就没有真正的教育。", "author": "马卡连柯"},
    {"quote_text": "教师的爱是滴滴甘露，即使枯萎的心灵也能苏醒。", "author": "巴甫连柯"}, # (需要注意，此处的“巴甫连柯”相对不如苏霍姆林斯基等人著名，但该句名言确有流传并归于此名下)
    {"quote_text": "教育者，非为已往，非为现在，乃专为将来。", "author": "蔡元培"},
    {"quote_text": "要尊重儿童的感情，也要引导儿童的感情。", "author": "鲁迅"},
    {"quote_text": "德者，才之帅也；才者，德之资也。", "author": "司马光"},
    {"quote_text": "智育是思想的提琴，而德育是操纵那提琴的弓。", "author": "斯特恩"}, # (劳伦斯·斯特恩，英国作家)
    {"quote_text": "教育的目的是培养身心和谐发展的人。", "author": "马卡连柯"},
    {"quote_text": "为人师表，以德立身。", "author": ""},
    {"quote_text": "育人先育德，正人先正己。", "author": ""},
    {"quote_text": "关爱是最好的教育。", "author": ""},
    {"quote_text": "学生是学习的主人，教师是学习的引路人。", "author": "叶圣陶"},
    {"quote_text": "教育贵在启发，而非灌输。", "author": "第斯多惠"}, # (阿道夫·第斯多惠，德国教育家)
    {"quote_text": "教育的艺术在于唤醒和激励。", "author": ""},
    {"quote_text": "以爱动其心，以理服其人。", "author": ""},
    {"quote_text": "中职教育重在培养学生的职业道德和职业技能。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "德技并修，知行合一。", "author": ""}, # (现代职业教育常用语，“知行合一”源自王阳明，但组合使用是现代提法)
    {"quote_text": "技术是工具，道德是导航。", "author": ""},
    {"quote_text": "中职学生更需要关爱和引导。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "用爱心点亮学生的未来。", "author": ""},
    {"quote_text": "中职德育要贴近学生实际，注重实践。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "把德育融入到职业技能教学中。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "以德立学，以技立业。", "author": ""},
    {"quote_text": "中职教育是为学生终身发展奠基。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "关爱每一个学生，成就每一个梦想。", "author": ""},
    {"quote_text": "教育不是注满一桶水，而是点燃一把火。", "author": "叶芝"}, # (常归于叶芝，也有说法源自普鲁塔克)
    {"quote_text": "教师是人类灵魂的工程师。", "author": "斯大林"}, # (此说法流传甚广，常归于斯大林)
    {"quote_text": "热爱学生是教师生活中最主要的东西。", "author": "苏霍姆林斯基"},
    {"quote_text": "教育的最高境界是爱。", "author": ""},
    {"quote_text": "尊重学生的个性，促进学生的全面发展。", "author": ""},
    {"quote_text": "好的先生不是教书，不是教学生，乃是教学生学。", "author": "陶行知"},
    {"quote_text": "千教万教，教人求真；千学万学，学做真人。", "author": "陶行知"},
    {"quote_text": "捧着一颗心来，不带半根草去。", "author": "陶行知"},
    {"quote_text": "爱满天下，德行千古。", "author": ""},
    {"quote_text": "以身作则，为人师表。", "author": ""},
    {"quote_text": "没有惩罚的教育是不完整的教育。", "author": "马卡连柯"},
    {"quote_text": "德育犹如树木的根，只有根深才能叶茂。", "author": ""},
    {"quote_text": "中职教育，不只是技能的传授，更是人格的塑造。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "用欣赏的眼光看待学生，用宽容的心态对待学生。", "author": ""},
    {"quote_text": "教育的真谛在于启发学生的内在力量。", "author": ""},
    {"quote_text": "关爱是打开学生心灵的钥匙。", "author": ""},
    {"quote_text": "中职德育，要注重学生的职业理想和职业责任感培养。", "author": ""}, # (针对中职教育的理念性表述)
    {"quote_text": "让每一个学生都找到自己的闪光点。", "author": ""},
    {"quote_text": "以爱为舟，载学生驶向成功的彼岸。", "author": ""},
    {"quote_text": "中职教育，是为学生的幸福人生奠基。", "author": ""} # (针对中职教育的理念性表述)
]
    response_data.update(random_quote=random.choice(quotes))
    return render_template('home.html', **response_data)


@app.route('/weiji_show', methods=['GET'])
@required_role('super_admin')
@check_redis_session  # 重启后重定向login
def weiji_show():
    # 加载 session
    response_data = load_session()

    try:
        # 加载数据
        department = response_data['user_department']
        data = load_data(department)

        # 获取搜索关键词
        search_term = request.args.get('search', '').strip().lower()

        # 根据搜索词过滤数据
        if search_term:
            data = data[
                data.apply(lambda row: any(
                    str(item).lower().find(search_term) != -1 for item in row), axis=1
                           )
            ]
        per_page = int(request.args.get('perPage', 10))  # 获取每页显示数量
        page = int(request.args.get('page', 1))  # 获取当前页

        # 计算分页信息
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        page_data = data.values.tolist()[start_index:end_index]

        total_records = len(data)
        total_pages = math.ceil(total_records / per_page)

        # 返回分页数据
        response_data.update({
            'records': page_data,
            'total': total_records,
            'totalPages': total_pages,
            'currentPage': page,
            'perPage': per_page,
            'weiji_types': WEIJI_TYPES,
        })

        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify(response_data)

        return render_template('weiji_show.html', **response_data)

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        return jsonify(error="数据文件未找到，请联系管理员。"), 500

    except pd.errors.EmptyDataError as e:
        logger.error(f"Empty data: {e}")
        return jsonify(error="数据文件为空，请联系管理员。"), 500

    except Exception as e:
        logger.exception("Error in weiji_show")  # 记录完整的堆栈信息
        return jsonify(error="服务器内部错误，请稍后再试。"), 500  # 返回 JSON 错误响应，并设置 HTTP 状态码


@app.route('/get_record', methods=['GET'])
@check_redis_session # 重启后重定向login
def get_record():

    # 获取选择的数据ID
    record_id = request.args.get('record_id')

    # 加载session
    response_data = load_session()
    
    # 加载数据
    department = response_data['user_department']
    data = load_data(department)

    try:
        record_id = int(record_id) # 转换id为int类型
        # 关键：这里通过 iloc[:, 0] 比较 ID，然后取 .values.tolist()[0]
        # 这会返回一个包含该行所有列数据的 *列表*，顺序与 CSV 文件中的列顺序一致
        record = data[data.iloc[:, 0] == record_id].values.tolist()[0]
        logger.debug(f"查找record_id={record_id}的记录，结果为：{record}")
        return jsonify(record) # 返回这个列表

    except (IndexError, TypeError):
        logger.error(f"找不到id={record_id}的记录")
        # 如果找不到记录，这里会返回 '找不到记录'，状态码是 200 OK，
        # 前端 fetchData 会认为成功，但在填充数据时 data 会是 {message: '找不到记录'}，导致JS错误
        # 应该返回错误状态码
        return jsonify(error='找不到记录'), 404 # 返回 404 Not Found
    except Exception as e: # 捕获其他潜在错误
        logger.exception(f"获取记录时发生错误, id={record_id}")
        return jsonify(error='服务器内部错误'), 500


@app.route('/update_record', methods=['POST'])
@check_redis_session  # 重启后重定向login
def update_record():
    record_id = request.args.get('record_id')
    updated_data = request.get_json()
    logger.debug(f"接收到id={record_id}的更新请求，数据为：{updated_data}")

    # 加载session
    response_data = load_session()
    
    # 加载数据
    department = response_data['user_department']
    data = load_data(department)

    selected_types_list = updated_data.get('type', []) # Get the list from JSON, default to empty list
    types_string = ','.join(selected_types_list)     # Join the list into a comma-separated string

    cols = data.columns
    try:
        record_id = int(record_id)  # 转换id为int类型
        data.loc[data.iloc[:, 0] == record_id, [
            cols[1],
            cols[8],
            cols[2],
            cols[3],
            cols[4],
            cols[5],
            cols[6],
            cols[7],
            cols[9],
        ]] = [
            updated_data['name'],
            updated_data['department'],
            updated_data['className'],
            updated_data['level'],
            updated_data['date'],
            # updated_data['type'],
            types_string,
            updated_data['reason'],
            updated_data['revoke'],
            updated_data['byteacher'],
        ]
        # 根据系部修改数据表
        if department == '经艺系':
            file = DATA_FILE_JINGYI
        elif department == '智能制造系':
            file = DATA_FILE_ZHIZAO
        data.to_csv(file, index=False)

        logger.info(f"id={record_id}的数据更新成功，更新后的数据为：{updated_data}")
        return jsonify(message='修改成功')
    except (IndexError, TypeError):
        logger.error(f"找不到id={record_id}的记录")
        return jsonify(message='修改失败')


@app.route('/delete_record', methods=['DELETE'])
@check_redis_session  # 重启后重定向login
def delete_record():
    record_id = request.args.get('record_id')
    logger.debug(f"接收到id={record_id}的删除请求")

    # 加载session
    response_data = load_session()
    
    # 加载数据
    department = response_data['user_department']
    data = load_data(department)

    try:
        record_id = int(record_id)  # 转换id为int类型
        
        # 将删除数据保存到 archive.csv 中以防误删
        deleted_data = data[data.iloc[:, 0] == record_id]  # 将其保存到archive.csv中
        archive_csv = pd.read_csv('data/archive.csv')
        archive_csv = pd.concat([archive_csv, deleted_data], ignore_index=False)
        archive_csv.to_csv('data/archive.csv', index=False)
        
        
        # 保存修改
        data = data[data.iloc[:, 0] != record_id]
        if department == '经艺系':
            file = DATA_FILE_JINGYI
        elif department == '智能制造系':
            file = DATA_FILE_ZHIZAO
        data.to_csv(file, index=False)
        logger.info(f"id={record_id}的数据删除成功")
        return jsonify(message='删除成功')
    except (IndexError, TypeError):
        logger.error(f"找不到id={record_id}的记录")
        return jsonify(message='删除失败')


@app.route('/weiji_add', methods=['GET', 'POST'])
@required_role('super_admin')
@check_redis_session  # 重启后重定向login
def weiji_add():
    response_data = load_session()
    today = datetime.date.today().strftime('%Y-%m-%d')
    response_data.update({'today': today})

    if request.method == 'POST':
        # --- POST 处理逻辑保持不变 ---
        try:
            # 加载数据
            department = response_data['user_department']
            df = load_data(department)

            new_id = df.iloc[:, 0].max() + 1 if not df.empty and pd.api.types.is_numeric_dtype(df.iloc[:, 0]) else 1

            selected_types = request.form.getlist('type')
            types_string = ','.join(selected_types) # 如果没有选择，会是空字符串 ''

            new_record = {
                '姓名': request.form['name'],
                '系部': request.form['department'],
                '班级': request.form['class'],
                '等级': request.form['level'],
                '日期': request.form['date'],
                # '类型': request.form['type'],
                '类型': types_string,
                # '经手人': request.form['byteacher'],
                '原因': request.form['reason'],
            }

            new_row_data = {
                df.columns[0]: new_id,
                df.columns[1]: new_record['姓名'],
                df.columns[2]: new_record['班级'],
                df.columns[3]: new_record['等级'],
                df.columns[4]: new_record['日期'],
                df.columns[5]: new_record['类型'],
                df.columns[6]: new_record['原因'],
                df.columns[7]: '', # 撤销信息默认为空字符串或 None
                df.columns[8]: new_record['系部'],
                # df.columns[8]: new_record['经手人'],
                df.columns[9]: session['username']
            }
            # 确保新行的列顺序和DataFrame一致
            new_row_df = pd.DataFrame([new_row_data], columns=df.columns)

            df = pd.concat([df, new_row_df], ignore_index=True)
            if department == '经艺系':
                file = DATA_FILE_JINGYI
            elif department == '智能制造系':
                file = DATA_FILE_ZHIZAO
            df.to_csv(file, index=False) # 指定编码

            # flash('记录添加成功!', 'success') # 可以用 flash 消息提示
            return redirect(url_for('weiji_show'))

        except Exception as e:
            logger.error(f"添加记录失败: {str(e)}", exc_info=True)
            # flash(f"添加失败: {str(e)}", "error") # 显示更具体的错误给用户（可选）
            # 重新渲染表单时，保留已输入的数据会更好，但这里简化处理
            return render_template('weiji_add.html',
                                   error="添加失败，请检查数据或联系管理员",
                                   departments=DEPARTMENTS,
                                   levels=['警告', '严重警告', '记过', '留校察看', '开除学籍'],
                                   weiji_types = WEIJI_TYPES,
                                   **response_data)

    # --- GET 请求处理保持不变 ---
    return render_template('weiji_add.html',
                           departments=DEPARTMENTS,
                           levels=['警告', '严重警告', '记过', '留校察看', '开除学籍'],
                           weiji_types = WEIJI_TYPES,
                           **response_data)

@app.route('/get_student_records', methods=['GET'])
@check_redis_session  # 同样需要登录态检查
def get_student_records():
    student_name = request.args.get('name')
    class_name = request.args.get('class_name') # 前端传来的参数名

    if not student_name or not class_name:
        # logger.warning("查询学生记录请求缺少姓名或班级参数")
        return jsonify({'records': [], 'message': '请输入姓名和班级进行查询'}), 400 # 返回 400 Bad Request

    try:
        # 加载session
        response_data = load_session()
        
        # 加载数据
        department = response_data['user_department']
        data = load_data(department)

        # 筛选记录 - 确保比较的是字符串类型（如果班级是数字，可能需要转换）
        # 假设 CSV 中的 '姓名' 和 '班级' 列都是字符串
        filtered_records = data[
            (data['姓名'] == student_name) &
            (data['班级'] == int(class_name)) # 显式转换为字符串以匹配
        ]

        # 按日期降序排序，最新的记录在前面
        if not filtered_records.empty and '日期' in filtered_records.columns:
            filtered_records = filtered_records.sort_values(by='日期', ascending=False)

        # 将筛选结果转换为字典列表以便 JSON 序列化
        records_list = filtered_records.to_dict('records')
        # logger.debug(f"为学生 {student_name} (班级 {class_name}) 查询到 {len(records_list)} 条记录")
        return jsonify({'records': records_list})

    except FileNotFoundError:
        logger.error(f"数据文件未找到:")
        return jsonify({'error': '数据文件未找到'}), 500
    except pd.errors.EmptyDataError:
         logger.error(f"数据文件为空:")
         return jsonify({'records': []}) # 文件为空也返回空列表
    except Exception as e:
        logger.exception(f"查询学生 {student_name} 记录时发生错误") # 记录完整错误信息
        return jsonify({'error': '查询记录时发生内部错误'}), 500


@app.route('/data_analysis')
@check_redis_session  # 重启后重定向login
def data_analysis():
    """
    数据分析路由，读取 weiji.csv 并统计数据，渲染到 data_analysis.html。
    """
    response_data = load_session()
    data_jyx = pd.read_csv('data/weiji_jingyixi.csv')
    data_zzx = pd.read_csv('data/weiji_zhizaoxi.csv')
    data = pd.concat([data_jyx, data_zzx])

    # 违纪违规总数
    total_count = len(data)

    # 经艺系违纪违规总数
    jingyixi_count = len(data[data['系部'] == '经艺系'])
    zhizaoxi_count = len(data[data['系部'] == '智能制造系'])
    
    # 已撤销人数
    revoked_count = len(data[data['撤销信息'].str.contains('已撤销', na=False)])

    # 未撤销人数
    not_revoked_count = total_count - revoked_count

    # 12个月
    year = '2025'
    months_format = [year + '/'  + str(n) + '/' for n in range(1, 13)]
    month_count = [len(data[data['日期'].str.contains(month, na=False)]) for month in months_format]

    response_data.update({
        'total_count': total_count,
        'jingyixi_count': jingyixi_count,
        'zhizaoxi_count': zhizaoxi_count,
        'revoked_count': revoked_count,
        'not_revoked_count': not_revoked_count,
        'new_month_count': month_count,
        'year': year
    })

    return render_template('data_analysis.html', **response_data)


@app.route('/data_plot_by_cate')
@check_redis_session
def data_plot_by_cate():
    labels = WEIJI_TYPES
    def count_type(labels, data_file) -> list:
        '''
        内置函数, 加载labels, 计算不同文件里的各标签数量, 返回values
        '''
        counts = []
        data = pd.read_csv(data_file)
        for label in labels:
            count = len(data[data['类型'] == label])
            counts.append(count)
        return counts
            
    jingyixi_counts = count_type(labels, 'data/weiji_jingyixi.csv')
    zhizaoxi_counts = count_type(labels, 'data/weiji_zhizaoxi.csv')
    
    response_data = load_session()
    data = {
    "labels": labels,
    "values_jingyixi": jingyixi_counts,
    "values_zhizaoxi": zhizaoxi_counts
    }
    response_data['data'] = data
    return render_template('data_plot_by_cate.html', **response_data)


@app.route('/data_plot_by_level')
@check_redis_session
def data_plot_by_level():
    labels = WEIJI_LEVELS[:-1]  # 去除开除学籍
    def count_type(labels, data_file) -> list:
        '''
        内置函数, 加载labels, 计算不同文件里的各标签数量, 返回values
        '''
        counts = []
        data = pd.read_csv(data_file)
        for label in labels:
            count = len(data[data['处分等级'] == label])
            counts.append(count)
        return counts
            
    jingyixi_counts = count_type(labels, 'data/weiji_jingyixi.csv')
    zhizaoxi_counts = count_type(labels, 'data/weiji_zhizaoxi.csv')
    
    response_data = load_session()
    data = {
    "labels": labels,
    "values_jingyixi": jingyixi_counts,
    "values_zhizaoxi": zhizaoxi_counts
    }
    response_data['data'] = data
    return render_template('data_plot_by_level.html', **response_data)


@app.route('/download_disposition/<int:record_id>')
@check_redis_session  # 重启后重定向login
def download_disposition(record_id):
    """
    根据记录 ID 生成并下载包含学生处分信息的 Word 文档。
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    # 加载session
    response_data = load_session()
    
    # 加载数据
    department = response_data['user_department']
    data = load_data(department)

    date = datetime.datetime.now().strftime("%Y年%m月%d日")

    try:
        record = data[data.iloc[:, 0] == record_id].values.tolist()[0]
        name = record[1]
        department = record[8]
        className = int(record[2])
        level = record[3]
        date = record[4]
        reason = record[6]
        prompt = ''
        if level in ['警告', '严重警告']:
            length = 3
        elif level == '记过':
            length = 6
        elif level == '留校察看':
            prompt = '4. 留校期间如再有违反校纪校规的行为，自愿退学\n'
            length = 12

        # 创建 Word 文档
        document = Document()

        # ===== 设置全局样式 =====
        style = document.styles['Normal']
        font = style.font
        font.name = '宋体'  # 主字体设置
        font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 显式设置东亚字体
        font.size = Pt(14)  # 三号字对应的磅值

        # ===== 文档标题 =====
        t = document.add_paragraph(style='Normal')
        t_run = t.add_run('学 生 处 分\n')
        t_run.font.name = '宋体'
        t_run.font.size = Pt(30)
        t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # ===== 正文内容 =====
        body = document.add_paragraph(style='Normal')  # 强制指定样式
        body_format = body.paragraph_format
        body_format.first_line_indent = Cm(0.74)
        body.add_run(
            f"{name}，系我校{className}班学生。该生因{reason}，违反了《中学生日常行为规范》及《嘉善技师学院·嘉善县中等专业学校学生管理规定》的相关条款，造成不良影响。经学校政教处研究决定，给予{level}处分。")

        # ===== 处分要求段落 =====
        document.add_paragraph("处分要求：", style='Normal')
        requirements = document.add_paragraph(style='Normal')  # 保持基础样式
        for run_text in [f"1. 即日起进入考察期，考察期为{length}个月\n",
                         "2. 完成违纪考察鉴定记录本, 每月到班主任处提交思想汇报\n",
                         "3. 家长需积极配合学校教育管理\n",
                         f"{prompt}"]:
            req_run = requirements.add_run(run_text)
            req_run.font.name = '宋体'  # 显式设置列表字体


        # ===== 落款部分 =====
        # 学校信息（右对齐）
        school_info = document.add_paragraph()
        school_info.add_run(f"嘉善技师学院·嘉善县中等专业学校 {department}")
        school_info.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_para = document.add_paragraph(f"日期 : {date}\n")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # ===== 签字部分 =====
        sign = document.add_paragraph()
        sign.add_run("学生签字：")
        sign.add_run("\n\n")  # 签字留空
        sign.add_run("家长签字：")
        sign.add_run("\n\n")
        sign.add_run("班主任签字：")


        # ===== 设置页边距 =====
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(3.7)  # 上边距3.7cm
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)


        # 保存文档到内存
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # 返回文件
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{name}_处分决定.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except (IndexError, TypeError):
        logger.error(f"找不到id={record_id}的记录")
        return "找不到记录", 404  # 返回一个错误页面或消息


@app.route('/class_weiji_show', methods=['GET'])
# @required_role('super_admin')
@check_redis_session  # 重启后重定向login
def class_weiji_show():
    # load session
    response_data = load_session()

    try:
        # 加载数据
        department = response_data['user_department']
        data = load_data(department)

        # data['班级'] = data['班级'].astype(int)
        # data = data[data['班级'] == int(session.get('class'))]
        class_value = session.get('class')
        data = data.query("班级 == @class_value")

        # 获取搜索关键词
        search_term = request.args.get('search', '').strip().lower()

        # 根据搜索词过滤数据
        if search_term:
            data = data[
                data.apply(lambda row: any(
                    str(item).lower().find(search_term) != -1 for item in row), axis=1
                           )
            ]
        per_page = int(request.args.get('perPage', 10))  # 获取每页显示数量
        page = int(request.args.get('page', 1))  # 获取当前页

        # 计算分页信息
        start_index = (page - 1) * per_page
        end_index = start_index + per_page
        page_data = data.values.tolist()[start_index:end_index]

        total_records = len(data)
        total_pages = math.ceil(total_records / per_page)

        # 返回分页数据
        response_data.update({
            'records': page_data,
            'total': total_records,
            'totalPages': total_pages,
            'currentPage': page,
            'perPage': per_page,
            'myvar': f'{data.shape} {len(data)}'
        })

        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify(response_data)

        return render_template('class_weiji_show.html', **response_data)

    except FileNotFoundError as e:
        logger.error(f"File not found: {e}")
        return jsonify(error="数据文件未找到，请联系管理员。"), 500

    except pd.errors.EmptyDataError as e:
        logger.error(f"Empty data: {e}")
        return jsonify(error="数据文件为空，请联系管理员。"), 500

    except Exception as e:
        logger.exception("Error in weiji_show")  # 记录完整的堆栈信息
        return jsonify(error="服务器内部错误，请稍后再试。"), 500  # 返回 JSON 错误响应，并设置 HTTP 状态码


# 测试markdown渲染功能
@app.route('/test', methods=['GET'])
@check_redis_session
def test():
    response_data = load_session()
    
    with open("content.md", "r", encoding="utf-8") as f:  # 确保文件存在并正确编码
        markdown_content = f.read()

    # 将 Markdown 转换为 HTML
    response_data['content'] = markdown.markdown(markdown_content)
    # html_content = markdown.markdown(markdown_content)

    # 渲染模板并将 HTML 内容传递给它
    return render_template('test.html', **response_data)
   

if __name__ == '__main__':
    app.run(debug=True)